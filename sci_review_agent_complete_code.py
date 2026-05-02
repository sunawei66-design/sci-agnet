"""
SciReview-Agent: 面向 SCI 综述论文的智能评审与重构 Agent
=====================================================

功能定位
- 结构诊断：判断综述主线、章节功能、重复冗余、逻辑断点
- 文献匹配：判断参考文献覆盖、引用堆砌、未支撑论点、章节证据强弱
- 图表规划：生成原创框架图、机制图、路线图、表格清单及插入位置
- 审稿模拟：模拟主编、方法学审稿人、领域专家、语言编辑四类意见
- 重构建议：给出摘要、章节、图表、引用和投稿适配的修改路线
- Token plan：先分块、再分 Agent、最后汇总，降低长文档反复调用成本

运行方式一：命令行
    pip install openai pydantic python-docx pypdf pandas streamlit tiktoken
    export OPENAI_API_KEY="你的 API Key"  # Windows 可用 set OPENAI_API_KEY=你的 API Key
    python sci_review_agent_complete_code.py --manuscript paper.docx --refs refs.txt --journal "Building and Environment" --out outputs

运行方式二：Streamlit 网页界面
    streamlit run sci_review_agent_complete_code.py

说明
- 没有 OPENAI_API_KEY 时，会自动进入 heuristic 模式，输出规则诊断报告。
- 有 OPENAI_API_KEY 时，会调用 OpenAI Responses API 进行多 Agent 分析。
- 默认模型可通过 --model 修改，例如 --model gpt-5.2。
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import textwrap
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from pydantic import BaseModel, Field
except Exception:  # pragma: no cover
    BaseModel = object  # type: ignore
    Field = lambda default=None, **kwargs: default  # type: ignore

# Optional dependencies
try:
    import tiktoken
except Exception:  # pragma: no cover
    tiktoken = None

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None  # type: ignore


# ============================================================
# 1. 数据结构
# ============================================================

@dataclass
class DocumentChunk:
    chunk_id: int
    title_hint: str
    text: str
    char_count: int
    estimated_tokens: int


@dataclass
class TokenPlan:
    total_chars: int
    estimated_total_tokens: int
    chunk_count: int
    max_chunk_tokens: int
    agent_calls: Dict[str, int]
    notes: List[str] = field(default_factory=list)


@dataclass
class AgentResult:
    agent_name: str
    task_name: str
    output_markdown: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class ReviewConfig(BaseModel):
    manuscript_path: str = Field(..., description="论文正文路径，支持 txt/md/docx/pdf")
    refs_path: Optional[str] = Field(None, description="参考文献文件路径，支持 txt/md/docx/pdf，可为空")
    journal: str = Field("Target SCI journal", description="目标期刊")
    paper_type: str = Field("Critical review", description="论文类型，如 Review, Critical review, Perspective")
    research_field: str = Field("Building science / energy / environment", description="研究领域")
    model: str = Field("gpt-5.2", description="OpenAI 模型名称")
    out_dir: str = Field("outputs", description="输出目录")
    max_chunk_tokens: int = Field(6000, description="单个正文分块的最大估计 token 数")
    language: str = Field("Chinese", description="输出语言，Chinese 或 English")
    temperature: float = Field(0.2, description="生成稳定性参数")
    use_api: bool = Field(True, description="是否调用 OpenAI API")


# ============================================================
# 2. 文件读取
# ============================================================

class FileReader:
    """读取论文与参考文献文件。"""

    @staticmethod
    def read_file(path: str) -> str:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = p.suffix.lower()
        if suffix in [".txt", ".md", ".rst"]:
            return p.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            return FileReader._read_docx(p)
        if suffix == ".pdf":
            return FileReader._read_pdf(p)
        raise ValueError(f"暂不支持该文件格式: {suffix}. 请使用 txt, md, docx 或 pdf。")

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            import docx  # python-docx
        except Exception as exc:
            raise ImportError("读取 docx 需要安装 python-docx: pip install python-docx") from exc

        doc = docx.Document(str(path))
        parts: List[str] = []
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                parts.append(txt)
        # 读取表格内容，避免表格完全丢失
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise ImportError("读取 pdf 需要安装 pypdf: pip install pypdf") from exc

        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                pages.append(f"\n\n[PDF Page {i + 1}]\n" + (page.extract_text() or ""))
            except Exception:
                pages.append(f"\n\n[PDF Page {i + 1}]\n[无法提取本页文本]")
        return "\n".join(pages)


# ============================================================
# 3. 文本预处理与 Token Plan
# ============================================================

class TextUtils:
    @staticmethod
    def normalize_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def estimate_tokens(text: str, model: str = "gpt-5.2") -> int:
        if tiktoken is not None:
            try:
                enc = tiktoken.encoding_for_model(model)
                return len(enc.encode(text))
            except Exception:
                pass
        # 中英文混合粗估：英文约 4 字符/token，中文约 1.5 字符/token
        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
        other_chars = max(len(text) - chinese_chars, 0)
        return int(chinese_chars / 1.5 + other_chars / 4) + 1

    @staticmethod
    def extract_heading_hint(text: str) -> str:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        for line in lines[:10]:
            if len(line) < 120 and re.match(r"^((\d+\.?)+\s+|[A-Z][A-Za-z\s]{3,}|摘要|Abstract|Introduction|Conclusion|结论)", line):
                return line[:100]
        return lines[0][:100] if lines else "Untitled chunk"

    @staticmethod
    def chunk_text(text: str, max_tokens: int, model: str) -> List[DocumentChunk]:
        """按标题和段落进行稳定分块，避免硬切断章节。"""
        text = TextUtils.normalize_text(text)
        paragraphs = re.split(r"\n\s*\n", text)
        chunks: List[DocumentChunk] = []
        buf: List[str] = []
        chunk_id = 1

        def flush() -> None:
            nonlocal buf, chunk_id
            if not buf:
                return
            chunk_text_joined = "\n\n".join(buf).strip()
            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    title_hint=TextUtils.extract_heading_hint(chunk_text_joined),
                    text=chunk_text_joined,
                    char_count=len(chunk_text_joined),
                    estimated_tokens=TextUtils.estimate_tokens(chunk_text_joined, model),
                )
            )
            chunk_id += 1
            buf = []

        for para in paragraphs:
            candidate = "\n\n".join(buf + [para])
            if TextUtils.estimate_tokens(candidate, model) > max_tokens and buf:
                flush()
                buf.append(para)
            else:
                buf.append(para)
        flush()
        return chunks

    @staticmethod
    def clean_json_like(text: str) -> str:
        text = text.strip()
        text = re.sub(r"^```(?:json|markdown|md)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
        return text


# ============================================================
# 4. 规则诊断模块，无 API 时可运行
# ============================================================

class HeuristicAnalyzer:
    """不调用 LLM 的基础诊断，用于 API 不可用时的降级输出。"""

    SECTION_PAT = re.compile(r"(?m)^(\d+(?:\.\d+)*\s+[^\n]{3,}|Abstract|摘要|Introduction|Conclusion|References|参考文献)\s*$")

    @staticmethod
    def split_sections(text: str) -> List[Tuple[str, str]]:
        matches = list(HeuristicAnalyzer.SECTION_PAT.finditer(text))
        if not matches:
            return [("全文", text)]
        sections = []
        for i, m in enumerate(matches):
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            sections.append((m.group(1).strip(), text[start:end].strip()))
        return sections

    @staticmethod
    def citation_stats(text: str) -> Dict[str, Any]:
        citation_patterns = re.findall(r"\[(?:\d+(?:\s*[,-]\s*\d+)*)+\]", text)
        long_strings = [c for c in citation_patterns if len(re.findall(r"\d+", c)) >= 6]
        return {
            "citation_count_approx": len(citation_patterns),
            "long_citation_strings": long_strings[:20],
            "long_citation_string_count": len(long_strings),
        }

    @staticmethod
    def redundancy_flags(text: str) -> List[str]:
        flags = []
        phrases = [
            "具有重要意义", "越来越受到关注", "已有研究表明", "综上所述", "值得注意的是",
            "However", "In recent years", "It is worth noting", "plays an important role"
        ]
        for ph in phrases:
            count = text.count(ph)
            if count >= 5:
                flags.append(f"高频套话或重复表达：'{ph}' 出现 {count} 次")
        return flags

    @staticmethod
    def section_table(text: str) -> str:
        sections = HeuristicAnalyzer.split_sections(text)
        rows = ["| 章节 | 字符数 | 估计 token | 初步风险 |", "|---|---:|---:|---|"]
        for title, body in sections:
            tokens = TextUtils.estimate_tokens(body)
            risk = []
            if len(body) < 500:
                risk.append("过短，可能论证不足")
            if len(body) > 12000:
                risk.append("过长，建议拆分或压缩")
            if HeuristicAnalyzer.citation_stats(body)["long_citation_string_count"] > 2:
                risk.append("引用串偏长")
            rows.append(f"| {title[:60]} | {len(body)} | {tokens} | {'；'.join(risk) if risk else '正常'} |")
        return "\n".join(rows)

    @staticmethod
    def run_basic_report(text: str, refs: str, config: ReviewConfig) -> str:
        stats = HeuristicAnalyzer.citation_stats(text)
        flags = HeuristicAnalyzer.redundancy_flags(text)
        refs_count = len([l for l in refs.splitlines() if l.strip()]) if refs else 0
        return f"""
# SciReview-Agent 规则诊断报告

> 当前未调用 OpenAI API，以下为本地规则诊断结果。建议设置 OPENAI_API_KEY 后运行多 Agent 深度评审。

## 1. 基本信息

- 目标期刊：{config.journal}
- 论文类型：{config.paper_type}
- 研究领域：{config.research_field}
- 正文字数粗估：{len(text)} 字符
- 估计 token：{TextUtils.estimate_tokens(text, config.model)}
- 参考文献条目粗估：{refs_count}
- 引用标记数量粗估：{stats['citation_count_approx']}
- 长引用串数量：{stats['long_citation_string_count']}

## 2. 章节长度与风险

{HeuristicAnalyzer.section_table(text)}

## 3. 高频问题提示

{chr(10).join('- ' + f for f in flags) if flags else '- 未发现明显高频套话，但仍需进行语义层面评审。'}

## 4. 初步修改建议

1. 摘要按“背景压缩、核心缺口、方法边界、综合发现、贡献”五句式重写。
2. 引言末尾必须明确本文不是资料罗列，而是提出何种分析框架。
3. 每个二级标题应承担不同功能，避免“技术进展”“应用现状”“挑战”之间重复。
4. 超过 6 篇的长引用串应拆分为定义支撑、机制支撑和应用支撑。
5. 表格应从“文献清单表”转向“证据压缩表”，每张表对应一个综合判断。
6. 图件应覆盖总框架、机制链条、分类边界、评价指标和未来路线，而不是只做装饰性图。
""".strip()


# ============================================================
# 5. LLM Client
# ============================================================

class LLMClient:
    def __init__(self, model: str, temperature: float = 0.2, use_api: bool = True):
        self.model = model
        self.temperature = temperature
        self.use_api = use_api and bool(os.getenv("OPENAI_API_KEY")) and OpenAI is not None
        self.client = OpenAI() if self.use_api else None

    def complete(self, instructions: str, user_input: str) -> str:
        if not self.use_api or self.client is None:
            return "[LLM API 未启用，本部分需要设置 OPENAI_API_KEY 后生成。]"
        try:
            # Responses API. 某些新模型可能不接受 temperature，因此异常时降级不传 temperature。
            try:
                resp = self.client.responses.create(
                    model=self.model,
                    instructions=instructions,
                    input=user_input,
                    temperature=self.temperature,
                )
            except TypeError:
                resp = self.client.responses.create(
                    model=self.model,
                    instructions=instructions,
                    input=user_input,
                )
            return getattr(resp, "output_text", str(resp))
        except Exception as exc:
            return f"[LLM 调用失败: {exc}]"


# ============================================================
# 6. Agent 基类与五类 Agent
# ============================================================

class BaseAgent:
    name = "BaseAgent"

    def __init__(self, llm: LLMClient, config: ReviewConfig):
        self.llm = llm
        self.config = config

    def run_on_chunks(self, chunks: List[DocumentChunk], refs: str = "") -> AgentResult:
        partials = []
        for chunk in chunks:
            prompt = self.build_chunk_prompt(chunk, refs)
            output = self.llm.complete(self.instructions(), prompt)
            partials.append(f"\n\n## Chunk {chunk.chunk_id}: {chunk.title_hint}\n\n{output}")

        merged_input = "\n\n".join(partials)
        summary = self.llm.complete(self.instructions(), self.build_merge_prompt(merged_input))
        return AgentResult(
            agent_name=self.name,
            task_name=self.task_name(),
            output_markdown=summary,
            metadata={"chunk_count": len(chunks)},
        )

    def instructions(self) -> str:
        raise NotImplementedError

    def task_name(self) -> str:
        raise NotImplementedError

    def build_chunk_prompt(self, chunk: DocumentChunk, refs: str) -> str:
        raise NotImplementedError

    def build_merge_prompt(self, partials: str) -> str:
        return f"""
请合并以下分块评审结果，删除重复意见，保留最关键、最可执行、最符合 SCI 综述投稿的判断。

目标期刊：{self.config.journal}
论文类型：{self.config.paper_type}
研究领域：{self.config.research_field}
输出语言：{self.config.language}

分块结果：
{partials}
""".strip()


class StructureDiagnosisAgent(BaseAgent):
    name = "StructureDiagnosisAgent"

    def task_name(self) -> str:
        return "结构诊断与主线评估"

    def instructions(self) -> str:
        return """
你是 SCI 一区综述论文结构评审专家。你的任务不是润色，而是判断论文是否具备高水平综述所需的清晰主线、批判性框架、章节功能分工和逻辑闭环。
必须重点检查：
1. 核心科学问题是否明确
2. 章节是否围绕同一主线推进
3. 是否存在资料罗列、重复小节、概念空转
4. 是否从研究共识、主要分歧、方法局限和未来框架四个层次展开
5. 是否符合目标期刊的综述风格
输出必须具体，指出问题所在章节和修改动作。不要泛泛表扬。
""".strip()

    def build_chunk_prompt(self, chunk: DocumentChunk, refs: str) -> str:
        return f"""
请评审以下论文分块的结构质量。

分块编号：{chunk.chunk_id}
标题线索：{chunk.title_hint}
目标期刊：{self.config.journal}

请按以下结构输出：
A. 本分块承担的章节功能
B. 与全文主线的贴合度，评分 1-10
C. 主要逻辑问题，逐条说明
D. 重复或可压缩内容
E. 建议重构方式，给出可直接执行的动作

论文分块：
{chunk.text}
""".strip()

    def build_merge_prompt(self, partials: str) -> str:
        return f"""
以下是结构诊断 Agent 对各分块的评审结果。请合并为一份“全文结构诊断报告”。

必须输出：
1. 全文一句话主线判断
2. 章节功能矩阵，使用 markdown 表格，列为：章节或模块、当前功能、主要问题、重构建议、优先级
3. 最需要合并或删除的内容
4. 最需要增强的理论框架或批判性视角
5. 适合 {self.config.journal} 的重构目录建议
6. 500 字以内的总体判断

分块评审结果：
{partials}
""".strip()


class LiteratureMatchingAgent(BaseAgent):
    name = "LiteratureMatchingAgent"

    def task_name(self) -> str:
        return "文献匹配与证据链评估"

    def instructions(self) -> str:
        return """
你是 SCI 综述论文参考文献和证据链审查专家。你需要判断正文论点是否被文献充分支撑，是否存在引用堆砌、引用错位、重要论点缺证据、参考文献使用不均、综述型文献与原始研究文献比例不当等问题。
重点不是核验每条文献真假，而是检查“论点—证据—章节功能”的匹配关系。
输出必须可执行，避免泛泛建议。
""".strip()

    def build_chunk_prompt(self, chunk: DocumentChunk, refs: str) -> str:
        refs_excerpt = refs[:12000] if refs else "[未提供参考文献目录]"
        return f"""
请检查以下论文分块的文献使用与证据链质量。

目标期刊：{self.config.journal}
标题线索：{chunk.title_hint}

请输出：
A. 本分块的核心论点
B. 每个核心论点所需的文献类型，区分定义文献、机制文献、方法文献、案例文献、综述文献
C. 当前引用方式的风险，例如长引用串、引用后置、证据不足、观点大于证据
D. 需要补强或拆分引用的位置
E. 建议插入的“证据链句式”模板

参考文献目录节选：
{refs_excerpt}

论文分块：
{chunk.text}
""".strip()

    def build_merge_prompt(self, partials: str) -> str:
        return f"""
请将以下分块结果合并为“文献匹配与证据链评估报告”。

必须输出：
1. 引用总体风险判断
2. 文献—章节匹配表，列为：章节模块、需要的文献类型、当前风险、修改动作
3. 长引用串处理原则
4. 哪些论点最需要原始研究支撑
5. 哪些位置可以用高质量综述替代过多散引
6. 可直接粘贴到论文方法部分的“文献筛选和证据分类”写法

分块结果：
{partials}
""".strip()


class FigurePlanningAgent(BaseAgent):
    name = "FigurePlanningAgent"

    def task_name(self) -> str:
        return "图表规划与插入位置设计"

    def instructions(self) -> str:
        return """
你是 SCI 综述论文图表设计专家，熟悉 Building and Environment, Applied Energy, RSER 等期刊的综述图表逻辑。
你的任务是把论文从“纯文字综述”转化为“框架清晰、证据压缩、图文互证”的投稿形态。
重点设计原创图、机制图、分类图、路线图、证据表，而不是装饰图。
输出要包括图表标题、图表目的、核心元素、插入位置、与上下文的承接句。
""".strip()

    def build_chunk_prompt(self, chunk: DocumentChunk, refs: str) -> str:
        return f"""
请为以下论文分块规划图表。

目标期刊：{self.config.journal}
标题线索：{chunk.title_hint}

请输出：
A. 本分块是否需要图或表，说明原因
B. 推荐图表类型，区分 framework figure, mechanism diagram, taxonomy map, evidence table, roadmap
C. 图表标题，英文优先
D. 图表核心结构，列出图中应包含的框、箭头、层级或表格字段
E. 最合适插入位置
F. 图前引出句和图后承接句

论文分块：
{chunk.text}
""".strip()

    def build_merge_prompt(self, partials: str) -> str:
        return f"""
请合并以下结果，形成“全文图表规划方案”。

必须输出：
1. 建议保留的 5-8 张核心图表清单
2. 建议删除或合并的表格类型
3. 每张图表的插入位置和上下文承接句
4. 至少 3 张原创框架图的 SVG 设计提示词，要求简洁、可编辑、横向 16:9、Times New Roman、大字号、适合 SCI 综述
5. 图表密度是否适合 {self.config.journal}

分块图表建议：
{partials}
""".strip()


class ReviewerSimulationAgent(BaseAgent):
    name = "ReviewerSimulationAgent"

    def task_name(self) -> str:
        return "多角色审稿模拟"

    def instructions(self) -> str:
        return """
你是 SCI 期刊投稿前评审模拟系统。请分别从主编、方法学审稿人、领域专家、语言编辑四个角色给出严格意见。
要求：
- 不要奉承
- 不要只说语言问题
- 必须指出是否可能被 desk reject
- 必须区分 Major issue, Moderate issue, Minor issue
- 每条意见都要有修改建议
""".strip()

    def build_chunk_prompt(self, chunk: DocumentChunk, refs: str) -> str:
        return f"""
请对以下论文分块进行投稿前审稿模拟。

目标期刊：{self.config.journal}
标题线索：{chunk.title_hint}

请按四类角色输出：
1. Handling Editor
2. Methodological Reviewer
3. Domain Reviewer
4. Language and Presentation Reviewer

每类角色均输出：主要问题、严重程度、修改建议。

论文分块：
{chunk.text}
""".strip()

    def build_merge_prompt(self, partials: str) -> str:
        return f"""
请将以下分块审稿意见合并为“投稿前模拟审稿报告”。

必须输出：
1. Desk reject 风险判断，低/中/高，并说明理由
2. Major revision 级别问题，不超过 8 条
3. Moderate revision 级别问题，不超过 10 条
4. Minor revision 级别问题，不超过 10 条
5. 主编意见草案
6. 作者修改路线，按 3 天、7 天、14 天分阶段

分块审稿意见：
{partials}
""".strip()


class ReconstructionAgent(BaseAgent):
    name = "ReconstructionAgent"

    def task_name(self) -> str:
        return "论文重构与压缩路线"

    def instructions(self) -> str:
        return """
你是 SCI 综述论文重构专家，擅长将冗长、资料堆砌型初稿改造成具有批判性框架和投稿竞争力的综述。
你需要给出可执行的重构方案，而不是直接重写全文。
重点包括摘要、引言、章节合并、表格压缩、图表设计、引用优化、结论升维。
""".strip()

    def build_chunk_prompt(self, chunk: DocumentChunk, refs: str) -> str:
        return f"""
请对以下论文分块提出重构和压缩建议。

目标期刊：{self.config.journal}
标题线索：{chunk.title_hint}

请输出：
A. 必须保留的核心内容
B. 可以删除的冗余内容类型
C. 可以合并到其他章节的内容
D. 建议改写的关键句式
E. 如何增强批判性和综述深度
F. 预计可压缩比例

论文分块：
{chunk.text}
""".strip()

    def build_merge_prompt(self, partials: str) -> str:
        return f"""
请将以下分块建议合并为“全文重构与压缩路线图”。

必须输出：
1. 推荐的新标题或副标题方向
2. 摘要重写框架，给出 5 句式模板
3. 引言重构方案
4. 正文章节合并方案
5. 表格压缩方案
6. 图件增强方案
7. 结论三步式写法，体现根本瓶颈、范式转移、统一框架
8. 10000 字以内压缩策略，按模块给出预计删减比例

分块重构建议：
{partials}
""".strip()


# ============================================================
# 7. Orchestrator
# ============================================================

class SciReviewOrchestrator:
    def __init__(self, config: ReviewConfig):
        self.config = config
        self.out_dir = Path(config.out_dir)
        self.out_dir.mkdir(parents=True, exist_ok=True)
        self.llm = LLMClient(config.model, config.temperature, config.use_api)

    def build_token_plan(self, manuscript: str, chunks: List[DocumentChunk]) -> TokenPlan:
        agent_calls = {
            "StructureDiagnosisAgent": len(chunks) + 1,
            "LiteratureMatchingAgent": len(chunks) + 1,
            "FigurePlanningAgent": len(chunks) + 1,
            "ReviewerSimulationAgent": len(chunks) + 1,
            "ReconstructionAgent": len(chunks) + 1,
        }
        total_tokens = TextUtils.estimate_tokens(manuscript, self.config.model)
        notes = [
            "先分块诊断，避免一次性处理长文导致遗漏。",
            "每个 Agent 先生成分块意见，再进行一次合并。",
            "文献目录只在文献匹配 Agent 中重点使用，降低无关 token 消耗。",
            "最终报告只汇总高优先级意见，避免重复输出。",
        ]
        return TokenPlan(
            total_chars=len(manuscript),
            estimated_total_tokens=total_tokens,
            chunk_count=len(chunks),
            max_chunk_tokens=self.config.max_chunk_tokens,
            agent_calls=agent_calls,
            notes=notes,
        )

    def save_markdown(self, filename: str, content: str) -> Path:
        p = self.out_dir / filename
        p.write_text(content, encoding="utf-8")
        return p

    def run(self) -> Dict[str, Path]:
        manuscript = TextUtils.normalize_text(FileReader.read_file(self.config.manuscript_path))
        refs = ""
        if self.config.refs_path:
            refs = TextUtils.normalize_text(FileReader.read_file(self.config.refs_path))

        chunks = TextUtils.chunk_text(manuscript, self.config.max_chunk_tokens, self.config.model)
        token_plan = self.build_token_plan(manuscript, chunks)
        saved: Dict[str, Path] = {}

        token_plan_md = self.render_token_plan(token_plan, chunks)
        saved["token_plan"] = self.save_markdown("00_token_plan.md", token_plan_md)

        if not self.llm.use_api:
            basic_report = HeuristicAnalyzer.run_basic_report(manuscript, refs, self.config)
            saved["heuristic_report"] = self.save_markdown("01_heuristic_report.md", basic_report)
            return saved

        agents: List[BaseAgent] = [
            StructureDiagnosisAgent(self.llm, self.config),
            LiteratureMatchingAgent(self.llm, self.config),
            FigurePlanningAgent(self.llm, self.config),
            ReviewerSimulationAgent(self.llm, self.config),
            ReconstructionAgent(self.llm, self.config),
        ]

        results: List[AgentResult] = []
        for idx, agent in enumerate(agents, start=1):
            result = agent.run_on_chunks(chunks, refs)
            results.append(result)
            safe_name = re.sub(r"[^a-zA-Z0-9_]+", "_", agent.name).lower()
            saved[safe_name] = self.save_markdown(f"{idx:02d}_{safe_name}.md", result.output_markdown)

        final_report = self.build_final_report(results, token_plan)
        saved["final_report"] = self.save_markdown("06_final_integrated_report.md", final_report)
        return saved

    def render_token_plan(self, plan: TokenPlan, chunks: List[DocumentChunk]) -> str:
        rows = ["| Chunk | 标题线索 | 字符数 | 估计 token |", "|---:|---|---:|---:|"]
        for c in chunks:
            rows.append(f"| {c.chunk_id} | {c.title_hint[:80]} | {c.char_count} | {c.estimated_tokens} |")
        calls = "\n".join([f"- {k}: {v} 次调用" for k, v in plan.agent_calls.items()])
        notes = "\n".join([f"- {n}" for n in plan.notes])
        return f"""
# Token Plan

- 总字符数：{plan.total_chars}
- 估计总 token：{plan.estimated_total_tokens}
- 分块数量：{plan.chunk_count}
- 单块最大 token：{plan.max_chunk_tokens}

## Agent 调用计划

{calls}

## 分块清单

{chr(10).join(rows)}

## Token 节省策略

{notes}
""".strip()

    def build_final_report(self, results: List[AgentResult], token_plan: TokenPlan) -> str:
        parts = [
            f"# SciReview-Agent 综合评审与重构报告\n",
            f"- 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"- 目标期刊：{self.config.journal}",
            f"- 论文类型：{self.config.paper_type}",
            f"- 研究领域：{self.config.research_field}",
            f"- 分块数量：{token_plan.chunk_count}",
            "\n---\n",
        ]
        for r in results:
            parts.append(f"\n# {r.task_name}\n\n{r.output_markdown}\n")
        parts.append("\n---\n# 最终执行清单\n")
        parts.append("""
| 优先级 | 修改任务 | 目标 |
|---|---|---|
| P0 | 明确全文唯一主线 | 避免章节并列堆砌 |
| P0 | 合并重复章节和重复表格 | 控制字数并增强线性叙事 |
| P1 | 重写摘要和引言末段 | 建立 gap-framework-contribution 链条 |
| P1 | 重构引用方式 | 从引用堆砌转为证据链支撑 |
| P1 | 增加原创框架图和机制图 | 提升综述识别度 |
| P2 | 模拟审稿意见逐条回应 | 降低 desk reject 和 major revision 风险 |
""".strip())
        return "\n".join(parts)


# ============================================================
# 8. Streamlit UI
# ============================================================

def run_streamlit_app() -> None:
    try:
        import streamlit as st
    except Exception:
        print("请先安装 Streamlit: pip install streamlit")
        return

    st.set_page_config(page_title="SciReview-Agent", layout="wide")
    st.title("SciReview-Agent：SCI 综述论文智能评审与重构 Agent")
    st.caption("结构诊断 · 文献匹配 · 图表规划 · 审稿模拟 · 重构路线 · Token plan")

    with st.sidebar:
        st.header("配置")
        journal = st.text_input("目标期刊", value="Building and Environment")
        paper_type = st.selectbox("论文类型", ["Critical review", "Review", "Perspective", "Systematic review"])
        research_field = st.text_input("研究领域", value="Building science / energy / environment")
        model = st.text_input("模型", value="gpt-5.2")
        max_chunk_tokens = st.slider("单块最大估计 token", 2000, 12000, 6000, step=500)
        language = st.selectbox("输出语言", ["Chinese", "English"], index=0)
        use_api = st.checkbox("调用 OpenAI API", value=True)

    manuscript_file = st.file_uploader("上传论文正文，支持 txt/md/docx/pdf", type=["txt", "md", "docx", "pdf"])
    refs_file = st.file_uploader("上传参考文献目录，可选，支持 txt/md/docx/pdf", type=["txt", "md", "docx", "pdf"])

    if st.button("开始评审", type="primary"):
        if manuscript_file is None:
            st.error("请先上传论文正文。")
            return

        tmp_dir = Path("streamlit_tmp")
        tmp_dir.mkdir(exist_ok=True)
        manuscript_path = tmp_dir / manuscript_file.name
        manuscript_path.write_bytes(manuscript_file.getbuffer())

        refs_path = None
        if refs_file is not None:
            refs_path_p = tmp_dir / refs_file.name
            refs_path_p.write_bytes(refs_file.getbuffer())
            refs_path = str(refs_path_p)

        cfg = ReviewConfig(
            manuscript_path=str(manuscript_path),
            refs_path=refs_path,
            journal=journal,
            paper_type=paper_type,
            research_field=research_field,
            model=model,
            max_chunk_tokens=max_chunk_tokens,
            language=language,
            use_api=use_api,
            out_dir="streamlit_outputs",
        )
        with st.spinner("正在运行多 Agent 评审，请等待当前页面完成输出。"):
            orchestrator = SciReviewOrchestrator(cfg)
            saved = orchestrator.run()

        st.success("评审完成。")
        for name, path in saved.items():
            content = path.read_text(encoding="utf-8")
            with st.expander(f"{name}: {path.name}", expanded=(name == "final_report")):
                st.markdown(content)
                st.download_button(
                    label=f"下载 {path.name}",
                    data=content,
                    file_name=path.name,
                    mime="text/markdown",
                )


# ============================================================
# 9. CLI
# ============================================================

def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="SciReview-Agent: SCI 综述论文智能评审与重构 Agent")
    parser.add_argument("--manuscript", required=False, help="论文正文路径，支持 txt/md/docx/pdf")
    parser.add_argument("--refs", required=False, help="参考文献目录路径，支持 txt/md/docx/pdf")
    parser.add_argument("--journal", default="Building and Environment", help="目标期刊")
    parser.add_argument("--paper-type", default="Critical review", help="论文类型")
    parser.add_argument("--field", default="Building science / energy / environment", help="研究领域")
    parser.add_argument("--model", default="gpt-5.2", help="OpenAI 模型")
    parser.add_argument("--out", default="outputs", help="输出目录")
    parser.add_argument("--max-chunk-tokens", type=int, default=6000, help="单块最大估计 token")
    parser.add_argument("--language", default="Chinese", choices=["Chinese", "English"], help="输出语言")
    parser.add_argument("--no-api", action="store_true", help="不调用 API，只做规则诊断")
    return parser.parse_args(argv)


def main(argv: Optional[List[str]] = None) -> None:
    # Streamlit 运行时不走 CLI
    if "streamlit" in sys.modules:
        run_streamlit_app()
        return

    args = parse_args(argv)
    if not args.manuscript:
        print("请提供 --manuscript，或使用 streamlit run sci_review_agent_complete_code.py 启动网页界面。")
        sys.exit(1)

    cfg = ReviewConfig(
        manuscript_path=args.manuscript,
        refs_path=args.refs,
        journal=args.journal,
        paper_type=args.paper_type,
        research_field=args.field,
        model=args.model,
        out_dir=args.out,
        max_chunk_tokens=args.max_chunk_tokens,
        language=args.language,
        use_api=not args.no_api,
    )
    orchestrator = SciReviewOrchestrator(cfg)
    saved = orchestrator.run()

    print("\n完成。输出文件：")
    for key, path in saved.items():
        print(f"- {key}: {path}")


if __name__ == "__main__":
    main()
